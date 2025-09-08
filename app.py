# Council Grant Finder & Auto‑Writer — Streamlit MVP (v0.1)
# -------------------------------------------------------
# One‑week build focused on councils (AUS). Features:
# - Scrape VIC Grants Gateway + GrantConnect (basic web scraping)
# - Council profile + keyword matching + relevance scoring
# - AI grant application draft (OpenAI)
# - Export to DOCX / PDF
# - Weekly alert preview (email text preview)
# - Demo data fallback to keep the app usable during development
# -------------------------------------------------------

import os
import io
import re
import json
import time
import textwrap
from datetime import datetime, timedelta
from typing import List, Dict, Any

import pandas as pd
import requests
from bs4 import BeautifulSoup

import streamlit as st
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# Optional OpenAI (draft writer)
try:
    from openai import OpenAI
    OPENAI_ENABLED = True
except Exception:
    OPENAI_ENABLED = False

# ----------------------------
# App Config
# ----------------------------
st.set_page_config(page_title="Council Grant Finder", page_icon="💰", layout="wide")

APP_BRAND = {
    "name": "Council Grant Finder",
    "tagline": "Never miss a grant again — find, match, and draft in minutes.",
    "primary_colour": "#0F766E",  # teal 700
}

DEMO_COUNCILS = [
    {"name": "Wyndham City Council", "state": "VIC", "population": 300000, "priorities": ["youth", "waste", "roads", "community"]},
    {"name": "City of Melbourne", "state": "VIC", "population": 150000, "priorities": ["innovation", "culture", "climate"]},
    {"name": "City of Yarra", "state": "VIC", "population": 100000, "priorities": ["arts", "cycling", "heritage"]},
]

CACHE_FILE = "grants.json"
USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0 Safari/537.36"
)

# ----------------------------
# Utils
# ----------------------------

def load_grants_cache() -> pd.DataFrame:
    if os.path.exists(CACHE_FILE):
        try:
            data = json.load(open(CACHE_FILE, "r"))
            return pd.DataFrame(data)
        except Exception:
            pass
    return pd.DataFrame(columns=[
        "source", "title", "amount", "deadline", "category", "eligibility", "state",
        "url", "posted_at"
    ])


def save_grants_cache(df: pd.DataFrame) -> None:
    try:
        # NaT/NaN safe conversion
        data = json.loads(df.to_json(orient="records", date_format="iso"))
        json.dump(data, open(CACHE_FILE, "w"), indent=2)
    except Exception as e:
        st.warning(f"Could not save grants cache: {e}")


def parse_money(text: str) -> str:
    if not text:
        return ""
    # crude money detection
    m = re.findall(r"\$\s?([0-9,.]+[kKmM]?)", text)
    return f"${m[0]}" if m else ""


def parse_deadline(text: str) -> str:
    if not text:
        return ""
    # try to find a date-like pattern
    m = re.search(r"(\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4})", text)
    if m:
        return m.group(1)
    m2 = re.search(r"(\d{4}-\d{2}-\d{2})", text)
    return m2.group(1) if m2 else ""


# ----------------------------
# Scrapers (basic, resilient; will gracefully fallback)
# ----------------------------

HEADERS = {"User-Agent": USER_AGENT}


def scrape_vic_grants_gateway(max_items: int = 20) -> List[Dict[str, Any]]:
    """Very light HTML scrape. If the site is JS-heavy or structure changes, returns empty list."""
    url = "https://www.vic.gov.au/grants"
    items = []
    try:
        resp = requests.get(url, headers=HEADERS, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        # Heuristic: find grant cards/links in main content
        for a in soup.select("a[href*='grants']"):
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if not title or len(title) < 6:
                continue
            # Fetch detail page (best-effort)
            grant_url = href if href.startswith("http") else f"https://www.vic.gov.au{href}"
            detail_text = ""
            try:
                d = requests.get(grant_url, headers=HEADERS, timeout=12)
                if d.ok:
                    detail_text = BeautifulSoup(d.text, "html.parser").get_text(" ")
            except Exception:
                pass
            items.append({
                "source": "VIC Grants Gateway",
                "title": title,
                "amount": parse_money(detail_text),
                "deadline": parse_deadline(detail_text),
                "category": "",
                "eligibility": detail_text[:500],
                "state": "VIC",
                "url": grant_url,
                "posted_at": datetime.utcnow().isoformat(),
            })
            if len(items) >= max_items:
                break
    except Exception:
        # swallow and return [] — we'll fallback later
        return []
    return items


def scrape_grantconnect(max_items: int = 20) -> List[Dict[str, Any]]:
    """GrantConnect (federal). We do a shallow scrape of the public listings page."""
    base = "https://www.grants.gov.au/"
    items = []
    try:
        resp = requests.get(base, headers=HEADERS, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for a in soup.select("a[href*='/government-grants']"):
            title = a.get_text(strip=True)
            href = a.get("href", "")
            if not title or len(title) < 6:
                continue
            url = href if href.startswith("http") else f"https://www.grants.gov.au{href}"
            detail_text = ""
            try:
                d = requests.get(url, headers=HEADERS, timeout=12)
                if d.ok:
                    detail_text = BeautifulSoup(d.text, "html.parser").get_text(" ")
            except Exception:
                pass
            items.append({
                "source": "GrantConnect (Federal)",
                "title": title,
                "amount": parse_money(detail_text),
                "deadline": parse_deadline(detail_text),
                "category": "",
                "eligibility": detail_text[:500],
                "state": "AU",
                "url": url,
                "posted_at": datetime.utcnow().isoformat(),
            })
            if len(items) >= max_items:
                break
    except Exception:
        return []
    return items


def fetch_all_sources() -> pd.DataFrame:
    vic = scrape_vic_grants_gateway()
    fed = scrape_grantconnect()
    data = vic + fed
    df = pd.DataFrame(data)
    if df.empty:
        df = demo_data()
    return df


# ----------------------------
# Demo Data (fallback)
# ----------------------------

def demo_data() -> pd.DataFrame:
    demo = [
        {
            "source": "VIC Grants Gateway",
            "title": "Community Youth Engagement Grant",
            "amount": "$250,000",
            "deadline": (datetime.utcnow() + timedelta(days=14)).strftime("%d %b %Y"),
            "category": "community, youth",
            "eligibility": "Open to VIC local governments supporting youth engagement initiatives.",
            "state": "VIC",
            "url": "https://www.vic.gov.au/grants/youth-demo",
            "posted_at": datetime.utcnow().isoformat(),
        },
        {
            "source": "GrantConnect (Federal)",
            "title": "Waste & Recycling Infrastructure Upgrade",
            "amount": "$1,200,000",
            "deadline": (datetime.utcnow() + timedelta(days=21)).strftime("%d %b %Y"),
            "category": "waste",
            "eligibility": "Australian local governments improving waste diversion.",
            "state": "AU",
            "url": "https://www.grants.gov.au/government-grants/waste-demo",
            "posted_at": datetime.utcnow().isoformat(),
        },
        {
            "source": "VIC Grants Gateway",
            "title": "Active Transport & Cycling Corridors",
            "amount": "$500,000",
            "deadline": (datetime.utcnow() + timedelta(days=10)).strftime("%d %b %Y"),
            "category": "transport, cycling",
            "eligibility": "VIC councils expanding active transport networks.",
            "state": "VIC",
            "url": "https://www.vic.gov.au/grants/transport-demo",
            "posted_at": datetime.utcnow().isoformat(),
        },
    ]
    return pd.DataFrame(demo)


# ----------------------------
# Matching & Scoring
# ----------------------------

def keyword_score(text: str, keywords: List[str]) -> int:
    if not text:
        return 0
    score = 0
    lower = text.lower()
    for kw in keywords:
        if kw.lower() in lower:
            score += 1
    return score


def match_and_rank(df: pd.DataFrame, council: Dict[str, Any], extra_keywords: List[str]) -> pd.DataFrame:
    if df.empty:
        return df
    # Filter by state relevance
    mask = (df["state"].isin([council.get("state", ""), "AU"]))
    filtered = df[mask].copy()

    # Compute relevance score
    council_kws = list(set(council.get("priorities", [])))
    all_kws = council_kws + extra_keywords

    scores = []
    for _, row in filtered.iterrows():
        text = " ".join([str(row.get("title", "")), str(row.get("category", "")), str(row.get("eligibility", ""))])
        s = keyword_score(text, all_kws)
        # Deadline urgency bonus
        deadline_txt = str(row.get("deadline", ""))
        urgent_bonus = 0
        try:
            dt = datetime.strptime(deadline_txt, "%d %b %Y")
            days_left = (dt - datetime.utcnow()).days
            if 0 <= days_left <= 14:
                urgent_bonus = 2
            elif 15 <= days_left <= 30:
                urgent_bonus = 1
        except Exception:
            pass
        scores.append(s + urgent_bonus)
    filtered["relevance_score"] = scores
    filtered.sort_values(["relevance_score", "deadline"], ascending=[False, True], inplace=True)
    return filtered.reset_index(drop=True)


# ----------------------------
# Draft Writer
# ----------------------------

def generate_application_draft(council: Dict[str, Any], grant: Dict[str, Any], goals: str = "") -> str:
    base_prompt = f"""
Write a persuasive 450–600 word grant application draft for {council['name']}.
Grant: {grant.get('title','')}
Source: {grant.get('source','')} | Amount: {grant.get('amount','')} | Deadline: {grant.get('deadline','')}
Council state: {council.get('state','')} | Population: {council.get('population','')}
Council priorities: {', '.join(council.get('priorities', []))}
Specific goals (if any): {goals}

Structure the response with:
- Project summary (2–3 sentences)
- Need & community impact
- Objectives & KPIs (use measurable targets)
- Delivery plan & timeline
- Budget outline & co-funding (if relevant)
- Evaluation & reporting
Use clear, confident, public‑sector appropriate language.
""".strip()

    if OPENAI_ENABLED and os.getenv("OPENAI_API_KEY"):
        try:
            client = OpenAI()
            completion = client.chat.completions.create(
                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role": "user", "content": base_prompt}],
                temperature=0.6,
            )
            return completion.choices[0].message.content.strip()
        except Exception as e:
            return f"[Local Draft Fallback]\n\n{local_draft(council, grant, goals)}\n\n(Note: OpenAI error: {e})"
    else:
        return local_draft(council, grant, goals)


def local_draft(council: Dict[str, Any], grant: Dict[str, Any], goals: str) -> str:
    # Simple templated draft if OpenAI is unavailable
    return textwrap.dedent(f"""
    PROJECT SUMMARY
    {council['name']} seeks support through the program “{grant.get('title','')}” to deliver a targeted initiative aligned to our priorities: {', '.join(council.get('priorities', []))}. The project will serve our community of approximately {council.get('population','N/A')} residents.

    NEED & IMPACT
    This funding will address a clearly identified local need and unlock measurable benefits across priority cohorts. The project aligns with state and federal policy directions and complements existing council strategies.

    OBJECTIVES & KPIs
    • Achieve 3–5 measurable outcomes within 12 months.
    • Reach 200+ participants with >80% satisfaction.
    • Reduce service gaps or improve access by 15–25%.

    DELIVERY PLAN & TIMELINE
    • Months 1–2: Mobilisation and procurement\n    • Months 3–9: Delivery and monitoring\n    • Months 10–12: Evaluation and reporting

    BUDGET & CO‑FUNDING
    The budget will allocate funds to staffing, materials, outreach, and evaluation. Council will contribute in‑kind support and governance oversight.

    EVALUATION & REPORTING
    We will track outputs, outcomes, and benefits using a simple dashboard and provide timely acquittals.
    """)


# ----------------------------
# Export Helpers
# ----------------------------

def export_docx(title: str, body: str) -> bytes:
    doc = Document()
    h = doc.add_heading(title, level=1)
    h.runs[0].font.size = Pt(16)
    for para in body.split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(10)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def export_pdf(title: str, body: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 50, height - 60
    c.setFont("Times-Roman", 16)
    c.drawString(x, y, title)
    y -= 30
    c.setFont("Times-Roman", 11)
    for para in body.split("\n\n"):
        for line in textwrap.wrap(para, width=95):
            if y < 60:
                c.showPage()
                y = height - 60
                c.setFont("Times-Roman", 11)
            c.drawString(x, y, line)
            y -= 14
        y -= 10
    c.showPage()
    c.save()
    return buf.getvalue()


# ----------------------------
# Email Preview (no provider required yet)
# ----------------------------

def weekly_email_preview(council: Dict[str, Any], df: pd.DataFrame) -> str:
    top = df.head(5)
    lines = [
        f"Subject: {council['name']}: {len(top)} new grants this week",
        "",
        f"Hi team, here are this week’s top matches for {council['name']}:",
        "",
    ]
    for _, r in top.iterrows():
        lines.append(f"• {r['title']} — {r.get('amount','')} — Deadline: {r.get('deadline','')}\n  {r['url']}")
    lines.append("")
    lines.append("Reply to generate an application draft for any of the above.")
    return "\n".join(lines)


# ----------------------------
# UI
# ----------------------------

st.markdown(f"""
<style>
:root {{ --brand: {APP_BRAND['primary_colour']}; }}
.stApp header {{ background: linear-gradient(90deg, var(--brand), #0891B2); }}
.block-container {{ padding-top: 1.5rem; }}
.key-metric {{ background:#f8fafc;border-radius:14px;padding:14px 16px;border:1px solid #e2e8f0; }}
</style>
<div style='display:flex;align-items:center;gap:12px;'>
  <h1 style='margin:0'>💰 {APP_BRAND['name']}</h1>
</div>
<p style='color:#334155;margin-top:-6px'>{APP_BRAND['tagline']}</p>
""")

with st.sidebar:
    st.subheader("Council Profile")
    council_names = [c["name"] for c in DEMO_COUNCILS]
    selected_name = st.selectbox("Choose council", council_names, index=0)
    council = next(c for c in DEMO_COUNCILS if c["name"] == selected_name)

    custom_kw = st.text_input("Extra keywords (comma‑separated)", "waste, youth, roads")
    extra_keywords = [k.strip() for k in custom_kw.split(",") if k.strip()]

    st.markdown("---")
    st.caption("Data refresh")
    do_fetch = st.button("🔄 Fetch latest grants", use_container_width=True)

    st.markdown("---")
    st.caption("Exports")
    enable_pdf = st.checkbox("Enable PDF export", value=True)
    enable_docx = st.checkbox("Enable DOCX export", value=True)

    st.markdown("---")
    st.caption("AI Settings")
    st.text("Uses OpenAI if OPENAI_API_KEY is set")

# Load or fetch data
if "grants_df" not in st.session_state:
    st.session_state.grants_df = load_grants_cache()

if do_fetch or st.session_state.grants_df.empty:
    with st.spinner("Fetching grants from sources (best‑effort)…"):
        df_new = fetch_all_sources()
        st.session_state.grants_df = df_new
        save_grants_cache(df_new)
        st.success(f"Loaded {len(df_new)} grants from sources.")

# Matching
df_all = st.session_state.grants_df.copy()
ranked = match_and_rank(df_all, council, extra_keywords)

colA, colB, colC, colD = st.columns(4)
colA.metric("Total grants in cache", len(df_all))
colB.metric("Matches for council", len(ranked))
soon = sum([1 for d in ranked.get("deadline", []) if isinstance(d, str) and d])
colC.metric("Grants with deadlines", soon)
colD.metric("OpenAI enabled", "Yes" if (OPENAI_ENABLED and os.getenv("OPENAI_API_KEY")) else "No")

st.markdown("### Top Matches")
if ranked.empty:
    st.info("No matches yet. Try refreshing data or broadening keywords.")
else:
    show_cols = ["source", "title", "amount", "deadline", "relevance_score", "url"]
    st.dataframe(ranked[show_cols], use_container_width=True, hide_index=True)

    st.markdown("#### Generate Application Draft")
    idx = st.number_input("Row # from table (0‑based)", min_value=0, max_value=max(0, len(ranked)-1), value=0, step=1)
    goals = st.text_area("Specific goals / notes (optional)", "")
    if st.button("✍️ Generate draft", type="primary"):
        row = ranked.iloc[int(idx)].to_dict()
        with st.spinner("Drafting application…"):
            draft = generate_application_draft(council, row, goals)
        st.success("Draft ready below.")
        st.session_state.current_draft = {"title": row.get("title", "Grant Application Draft"), "body": draft}

# Show draft & exports
if "current_draft" in st.session_state:
    st.markdown("---")
    st.markdown(f"### 📄 {st.session_state.current_draft['title']}")
    st.text_area("Draft (you can edit before export)", value=st.session_state.current_draft["body"], height=420, key="draft_text")

    exp_col1, exp_col2 = st.columns(2)
    with exp_col1:
        if enable_docx and st.button("⬇️ Download DOCX"):
            data = export_docx(st.session_state.current_draft["title"], st.session_state.draft_text)
            st.download_button(
                label="Save .docx",
                data=data,
                file_name="grant_draft.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    with exp_col2:
        if enable_pdf and st.button("⬇️ Download PDF"):
            data = export_pdf(st.session_state.current_draft["title"], st.session_state.draft_text)
            st.download_button(
                label="Save .pdf",
                data=data,
                file_name="grant_draft.pdf",
                mime="application/pdf",
            )

# Weekly email preview
st.markdown("---")
st.markdown("### ✉️ Weekly Alert Preview")
preview = weekly_email_preview(council, ranked)
st.code(preview, language="text")

# Footer / notes
st.markdown("---")
st.caption("This MVP performs best‑effort scraping of public listings and provides AI‑assisted drafting. Always review eligibility and guidelines on the official grant page before applying.")
